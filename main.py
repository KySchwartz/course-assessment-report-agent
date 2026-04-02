import pandas as pd
from docx import Document
import os

def generate_reports_by_major(excel_file):
    # 1. Load Data
    df_students = pd.read_excel(excel_file, sheet_name='Grades')
    df_mapping = pd.read_excel(excel_file, sheet_name='Objectives')
    
    passing_grades = ['A', 'B', 'C']
    unique_majors = df_students['Major'].unique()

    for major in unique_majors:
        # Filter students by current major
        major_df = df_students[df_students['Major'] == major]
        num_students = len(major_df)
        
        # Initialize Word Doc for this Major
        doc = Document()
        doc.add_heading(f'Course Assessment Report: {major}', 0)
        
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        
        # Set Headers
        headers = ['COURSE OBJECTIVE', 'ACTIVITY', '# STUDENTS ASSESSED', 
                   '# MET', '# DID NOT MEET', '% MET', 'COMMENTS']
        for i, text in enumerate(headers):
            table.rows[0].cells[i].text = text

        # 2. Process each Objective from the Mapping sheet
        for _, row in df_mapping.iterrows():
            obj_id = row['Objective ID']
            obj_desc = row['Objective Description']
            # Split the assignments string into a list
            linked_assignments = [a.strip() for a in str(row['Assignments']).split(',')]
            
            met_count = 0
            
            # Check each student in the major
            for _, student in major_df.iterrows():
                # A student meets the objective if they pass ALL linked assignments
                passed_all = True
                for assignment in linked_assignments:
                    if student[assignment] not in passing_grades:
                        passed_all = False
                        break
                
                if passed_all:
                    met_count += 1
            
            # Calculations
            not_met_count = num_students - met_count
            percent_met = (met_count / num_students * 100) if num_students > 0 else 0
            
            # 3. Add row to table
            row_cells = table.add_row().cells
            row_cells[0].text = f"{obj_id}:\n{obj_desc}"
            row_cells[1].text = ", ".join(linked_assignments)
            row_cells[2].text = str(num_students)
            row_cells[3].text = str(met_count)
            row_cells[4].text = str(not_met_count)
            row_cells[5].text = f"{percent_met:.0f}%"
            
            # Logic-based comment
            if percent_met >= 80:
                row_cells[6].text = "The most of students met the objective. No change needed"
            elif percent_met >= 70:
                row_cells[6].text = "Less than 80 percent of students passed. Minor adjustments possible."
            else:
                row_cells[6].text = "Significant number of students did not meet objective. Review required."

        # Save the file named by major
        output_name = f"Assessment_Report_{major}.docx"
        doc.save(output_name)
        print(f"Created report for {major}: {output_name}")

if __name__ == "__main__":
    file_path = "test_student_data.xlsx"
    if os.path.exists(file_path):
        generate_reports_by_major(file_path)
    else:
        print("Excel file not found. Please run your data generator script first.")